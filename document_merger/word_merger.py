"""
Word document merging functionality.
"""

import os
from pathlib import Path
from typing import List, Optional, Tuple
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from .toc_manager import TOCManager
from .utils import validate_file_path, is_docx, create_output_directory, extract_headings_from_docx
from .config import Config


class WordMerger:
    """Merge multiple Word documents into a single document with auto-refreshing TOC."""
    
    def __init__(self, config: Optional[Config] = None):
        """Initialize Word document merger.
        
        Args:
            config: Configuration object (optional)
        """
        self.config = config or Config()
        self.doc_files: List[str] = []
        self.master_doc: Optional[Document] = None
        self.toc_manager = TOCManager()
        self.page_breaks_between_docs = True
        self.preserve_styles = self.config.get('word.keep_styles', True)

    def add_document(self, file_path: str, add_page_break: bool = True) -> 'WordMerger':
        """Add a Word document to merge queue.
        
        Args:
            file_path: Path to DOCX file
            add_page_break: Whether to add page break before this document
            
        Returns:
            Self for method chaining
        """
        validate_file_path(file_path, '.docx')
        self.doc_files.append((file_path, add_page_break))
        return self

    def add_documents_batch(self, file_paths: List[str], 
                           add_page_breaks: bool = True) -> 'WordMerger':
        """Add multiple Word documents at once.
        
        Args:
            file_paths: List of DOCX file paths
            add_page_breaks: Whether to add page breaks between documents
            
        Returns:
            Self for method chaining
        """
        for file_path in file_paths:
            self.add_document(file_path, add_page_breaks)
        return self

    def merge(self, master_file: Optional[str] = None) -> Document:
        """Merge all documents into master document.
        
        Args:
            master_file: Optional path to master document to start with
            
        Returns:
            Merged Document object
        """
        if master_file:
            validate_file_path(master_file, '.docx')
            self.master_doc = Document(master_file)
        else:
            # Create new document
            self.master_doc = Document()
        
        # Merge each document
        for file_path, add_break in self.doc_files:
            self._merge_document(file_path, add_break)
        
        return self.master_doc

    def _merge_document(self, file_path: str, add_page_break: bool) -> None:
        """Merge a single document into master.
        
        Args:
            file_path: Path to document to merge
            add_page_break: Whether to add page break before this document
        """
        doc = Document(file_path)
        
        # Add page break if requested
        if add_page_break and self.master_doc.paragraphs:
            self.master_doc.add_page_break()
        
        # Copy paragraphs
        for para in doc.paragraphs:
            # Add new paragraph to master
            new_para = self.master_doc.add_paragraph()
            
            # Copy formatting
            new_para.style = para.style
            new_para.alignment = para.alignment
            
            # Copy runs (text with formatting)
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                self._copy_run_formatting(run, new_run)
        
        # Copy tables
        for table in doc.tables:
            self._copy_table(table)
        
        # Extract headings for TOC
        headings = extract_headings_from_docx(file_path)
        for heading_text, level in headings:
            self.toc_manager.add_entry(heading_text, 0, level)  # Page num will be updated

    def _copy_run_formatting(self, source_run, target_run) -> None:
        """Copy formatting from source run to target run.
        
        Args:
            source_run: Source run object
            target_run: Target run object
        """
        try:
            # Copy font properties
            if source_run.font.bold is not None:
                target_run.font.bold = source_run.font.bold
            if source_run.font.italic is not None:
                target_run.font.italic = source_run.font.italic
            if source_run.font.underline is not None:
                target_run.font.underline = source_run.font.underline
            if source_run.font.size is not None:
                target_run.font.size = source_run.font.size
            if source_run.font.color.rgb is not None:
                target_run.font.color.rgb = source_run.font.color.rgb
            if source_run.font.name is not None:
                target_run.font.name = source_run.font.name
        except Exception as e:
            print(f"Warning: Could not copy run formatting: {e}")

    def _copy_table(self, table) -> None:
        """Copy a table to master document.
        
        Args:
            table: Table object to copy
        """
        try:
            # Create new table with same dimensions
            new_table = self.master_doc.add_table(rows=len(table.rows), 
                                                   cols=len(table.columns))
            new_table.style = table.style
            
            # Copy cells
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    # Copy text
                    new_table.rows[i].cells[j].text = cell.text
                    
                    # Copy cell formatting if possible
                    try:
                        new_table.rows[i].cells[j].width = cell.width
                    except:
                        pass
        except Exception as e:
            print(f"Warning: Could not copy table: {e}")

    def generate_toc(self, max_depth: int = 3, insert_at_beginning: bool = True) -> str:
        """Generate and insert table of contents.
        
        Args:
            max_depth: Maximum heading depth to include
            insert_at_beginning: Whether to insert TOC at beginning of document
            
        Returns:
            Formatted TOC string
        """
        if not self.master_doc:
            self.merge()
        
        # Extract headings from merged document
        self._extract_headings_from_master()
        
        # Filter by max depth
        if max_depth < self.toc_manager.get_max_level():
            filtered_toc = self.toc_manager.filter_by_level(max_depth)
            entries = filtered_toc.get_entries()
        else:
            entries = self.toc_manager.get_entries()
        
        if insert_at_beginning:
            self._insert_toc_page(entries)
        
        return self.toc_manager.format_toc()

    def _extract_headings_from_master(self) -> None:
        """Extract headings from master document paragraphs."""
        self.toc_manager.clear()
        page_num = 1
        
        for para in self.master_doc.paragraphs:
            style_name = para.style.name if para.style else ''
            
            # Check if this is a heading style
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.split()[-1])
                    self.toc_manager.add_entry(para.text, page_num, level)
                except (ValueError, IndexError):
                    pass

    def _insert_toc_page(self, entries: List) -> None:
        """Insert TOC page at beginning of document.
        
        Args:
            entries: List of TOC entries
        """
        # Create TOC in Word by inserting at beginning
        # Get current first paragraph
        first_para = self.master_doc.paragraphs[0]._element
        parent = first_para.getparent()
        
        # Create new TOC paragraph
        try:
            from docx.oxml import OxmlElement
            
            # Insert title
            toc_title = self.master_doc.add_paragraph('Table of Contents', style='Heading 1')
            
            # Insert entries
            for entry in entries:
                indent_level = '  ' * (entry.level - 1)
                self.master_doc.add_paragraph(
                    f"{indent_level}{entry.title}",
                    style=f'List {entry.level}' if entry.level <= 3 else 'Normal'
                )
            
            # Add page break after TOC
            self.master_doc.add_page_break()
        except Exception as e:
            print(f"Warning: Could not insert TOC page: {e}")

    def refresh_toc(self, max_depth: int = 3) -> str:
        """Refresh/update the table of contents.
        
        Args:
            max_depth: Maximum heading depth
            
        Returns:
            Updated TOC string
        """
        return self.generate_toc(max_depth, insert_at_beginning=False)

    def update_toc_fields(self) -> None:
        """Update TOC fields in document."""
        if not self.master_doc:
            return
        
        # Word uses special field codes for TOC
        # This would need to be done through Word COM interface
        print("Note: Open the document in Word and press Ctrl+A then F9 to update TOC fields")

    def save(self, output_path: str) -> None:
        """Save merged document to file.
        
        Args:
            output_path: Path to output DOCX file
        """
        if not self.master_doc:
            self.merge()
        
        create_output_directory(output_path)
        self.master_doc.save(output_path)
        
        print(f"✓ Merged Word document saved to: {output_path}")
        print(f"  Total paragraphs: {len(self.master_doc.paragraphs)}")
        print(f"  Total documents merged: {len(self.doc_files)}")

    def get_document(self) -> Document:
        """Get the master document object.
        
        Returns:
            Master Document object
        """
        if not self.master_doc:
            self.merge()
        return self.master_doc

    def clear(self) -> None:
        """Clear all data and reset merger."""
        self.doc_files.clear()
        self.master_doc = None
        self.toc_manager.clear()

    def __repr__(self) -> str:
        """String representation."""
        return f"WordMerger({len(self.doc_files)} documents)"
