"""
Test data generator for document merger tests.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch


def create_sample_docx_files():
    """Create sample DOCX files for testing."""
    output_dir = Path('tests/sample_docs')
    output_dir.mkdir(parents=True, exist_ok=True)
    
    documents = [
        {
            'filename': 'document1.docx',
            'title': 'Introduction to Document Merger',
            'sections': [
                {
                    'heading': 'Chapter 1: Getting Started',
                    'content': 'This document demonstrates the document merger functionality.\n\nIt shows how to combine multiple documents automatically.'
                },
                {
                    'heading': 'Section 1.1: Basic Concepts',
                    'content': 'Document merging is a process of combining multiple files into a single document. This is useful for creating comprehensive reports, books, or presentations.'
                },
            ]
        },
        {
            'filename': 'document2.docx',
            'title': 'Main Content and Features',
            'sections': [
                {
                    'heading': 'Chapter 2: Core Features',
                    'content': 'The document merger supports various features including:\n- PDF merging\n- Word document merging\n- Automatic TOC generation\n- Cross-format merging'
                },
                {
                    'heading': 'Section 2.1: PDF Support',
                    'content': 'PDF files can be merged with bookmarks and outlines automatically extracted.'
                },
                {
                    'heading': 'Section 2.2: Word Support',
                    'content': 'Word documents are merged while preserving formatting and styles.'
                },
            ]
        },
        {
            'filename': 'document3.docx',
            'title': 'Conclusion and Summary',
            'sections': [
                {
                    'heading': 'Chapter 3: Advanced Features',
                    'content': 'Advanced features include:\n- Hybrid document merging (PDF + Word)\n- Automatic page numbering\n- Custom TOC styling\n- Bookmark hierarchy support'
                },
                {
                    'heading': 'Section 3.1: Best Practices',
                    'content': 'When merging documents, follow these best practices:\n1. Organize documents in logical order\n2. Use consistent heading styles\n3. Review the final output\n4. Verify all bookmarks and links'
                },
            ]
        },
    ]
    
    for doc_info in documents:
        doc = Document()
        
        # Add title
        title = doc.add_heading(doc_info['title'], level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add sections
        for section in doc_info['sections']:
            heading = doc.add_heading(section['heading'], level=1)
            content = doc.add_paragraph(section['content'])
            doc.add_paragraph()  # Empty paragraph for spacing
        
        # Save document
        output_path = output_dir / doc_info['filename']
        doc.save(output_path)
        print(f"✓ Created: {output_path}")


def create_sample_pdf_files():
    """Create sample PDF files for testing."""
    output_dir = Path('tests/sample_pdfs')
    output_dir.mkdir(parents=True, exist_ok=True)
    
    documents = [
        {
            'filename': 'document1.pdf',
            'title': 'PDF Document 1: Introduction',
            'content': [
                'Chapter 1: Getting Started',
                '',
                'This is a sample PDF document demonstrating the PDF merger functionality.',
                'PDF files can be merged with bookmarks and automatic TOC generation.',
                '',
                'Section 1.1: Basic Concepts',
                'Document merging combines multiple PDF files into one.',
            ]
        },
        {
            'filename': 'document2.pdf',
            'title': 'PDF Document 2: Features',
            'content': [
                'Chapter 2: Core Features',
                '',
                'The PDF merger supports:',
                '- Automatic bookmark extraction',
                '- Page number tracking',
                '- Hierarchical bookmark support',
                '- Custom TOC generation',
                '',
                'Section 2.1: Bookmarks',
                'Bookmarks are automatically extracted from PDF outlines.',
            ]
        },
        {
            'filename': 'document3.pdf',
            'title': 'PDF Document 3: Conclusion',
            'content': [
                'Chapter 3: Advanced Features',
                '',
                'Advanced capabilities include:',
                '- Cross-format merging (PDF + Word)',
                '- Automatic TOC insertion',
                '- Style preservation',
                '- Batch processing',
                '',
                'Section 3.1: Best Practices',
                'Organize documents logically and verify bookmarks after merging.',
            ]
        },
    ]
    
    for doc_info in documents:
        output_path = output_dir / doc_info['filename']
        c = canvas.Canvas(str(output_path), pagesize=letter)
        width, height = letter
        
        # Add title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(inch, height - inch, doc_info['title'])
        
        # Add content
        y = height - 1.5 * inch
        c.setFont("Helvetica", 10)
        
        for line in doc_info['content']:
            if line.startswith('Chapter') or line.startswith('Section'):
                c.setFont("Helvetica-Bold", 12)
            else:
                c.setFont("Helvetica", 10)
            
            c.drawString(inch, y, line)
            y -= 0.3 * inch
            
            if y < 0.5 * inch:
                c.showPage()
                y = height - inch
        
        c.save()
        print(f"✓ Created: {output_path}")


def main():
    """Generate all test files."""
    print("="*60)
    print("Test Document Generator")
    print("="*60)
    
    print("\n1. Creating sample Word documents...")
    create_sample_docx_files()
    
    print("\n2. Creating sample PDF documents...")
    create_sample_pdf_files()
    
    print("\n3. Creating output directory...")
    output_dir = Path('tests/output')
    output_dir.mkdir(parents=True, exist_ok=True)
    print(f"✓ Created: {output_dir}")
    
    print("\n" + "="*60)
    print("Test files generated successfully!")
    print("="*60)


if __name__ == '__main__':
    main()
